from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,  WD_LINE_SPACING

from states.research_state import ResearchState



def format_heading(heading, size, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """
    Applies consistent formatting to headings.
    """

    heading.alignment = alignment

    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = True


def add_paragraph(document, text):
    """
    Adds a justified paragraph with Times New Roman formatting.
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(8)

    run = paragraph.add_run(text)

    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    
def document_generator(state: ResearchState) -> dict:
    """
    Converts the final markdown report into a professionally
    formatted DOCX document stored inside a BytesIO buffer.
    """

    report = state["final_report_with_citation"]

    document = Document()



    section = document.sections[0]

    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = document.styles["Normal"]

    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(8)

    for line in report.splitlines():

        line = line.strip()

        if not line:
            document.add_paragraph()
            continue
        if line.startswith("#"):

            heading = document.add_heading(level=1)

            heading.add_run(line[2:])

            format_heading(
                heading,
                size=22,
                alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            )
        elif line.startswith("##"):

            heading = document.add_heading(level=2)

            heading.add_run(line[3:])

            format_heading(
                heading,
                size=18
            )
        elif line.startswith("###"):

            heading = document.add_heading(level=3)

            heading.add_run(line[4:])

            format_heading(
                heading,
                size=15
            )

        elif line.startswith("-"):

            paragraph = document.add_paragraph(style="List Bullet")

            paragraph.paragraph_format.line_spacing = 1.5

            run = paragraph.add_run(line[2:])

            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        elif (
            len(line) > 2
            and line[0].isdigit()
            and line[1] == "."
        ):

            paragraph = document.add_paragraph(style="List Number")

            paragraph.paragraph_format.line_spacing = 1.5

            run = paragraph.add_run(line[3:])

            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        else:

            add_paragraph(document, line)

 

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return {
        "docx_buffer": buffer
    }